from __future__ import annotations

import json
import math
from pathlib import Path
from typing import Iterable, Sequence

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from project_inventory import CATEGORY_ORDER, build_inventory


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "pdf"
SUPPORT_DIR = ROOT / "output" / "technical_report"
DOCX_OUT = OUTPUT_DIR / "DataVerse_AI_Complete_Technical_Implementation_Report_2026.docx"
ARCHITECTURE_IMAGE = ROOT / "output" / "architecture" / "dataverse-production-architecture.png"
MODEL_METADATA_PATH = ROOT / "models" / "model_metadata.json"

NAVY = "17365D"
BLUE = "2563EB"
PURPLE = "7C3AED"
GREEN = "059669"
INK = "172033"
SLATE = "475569"
LIGHT = "F4F7FB"
MID = "D9E2F2"
WHITE = "FFFFFF"
RED = "B42318"
AMBER = "B54708"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_keep_with_next(paragraph, enabled: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if enabled and keep is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not enabled and keep is not None:
        p_pr.remove(keep)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, size: float = 8.2) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def format_bytes(value: int) -> str:
    size = float(value)
    for unit in ("B", "KB", "MB", "GB"):
        if size < 1024 or unit == "GB":
            return f"{size:,.1f} {unit}" if unit != "B" else f"{int(size):,} B"
        size /= 1024
    return f"{int(value):,} B"


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Update this field to generate the table of contents."
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2, placeholder, fld_char3])


class ReportBuilder:
    def __init__(self) -> None:
        self.doc = Document()
        self.table_no = 0
        self.figure_no = 0
        self._setup_document()

    def _setup_document(self) -> None:
        section = self.doc.sections[0]
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.68)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
        section.header_distance = Inches(0.32)
        section.footer_distance = Inches(0.32)

        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(10)
        normal.font.color.rgb = RGBColor.from_string(INK)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.08

        heading_specs = {
            "Title": (29, NAVY, 0, 10),
            "Heading 1": (18, NAVY, 14, 7),
            "Heading 2": (13.5, BLUE, 10, 5),
            "Heading 3": (11, PURPLE, 7, 3),
        }
        for name, (size, color, before, after) in heading_specs.items():
            style = styles[name]
            style.font.name = "Aptos Display"
            style.font.bold = True
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        caption = styles["Caption"]
        caption.font.name = "Aptos"
        caption.font.size = Pt(8.5)
        caption.font.bold = True
        caption.font.color.rgb = RGBColor.from_string(SLATE)
        caption.paragraph_format.space_before = Pt(3)
        caption.paragraph_format.space_after = Pt(5)

        header = section.header.paragraphs[0]
        header.text = "DataVerse AI | Complete Current-Folder Technical Audit"
        header.style = styles["Header"]
        header.runs[0].font.name = "Aptos"
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = RGBColor.from_string(SLATE)
        header.runs[0].font.bold = True
        add_page_number(section.footer.paragraphs[0])

        settings = self.doc.settings._element
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)

    def page_break(self) -> None:
        self.doc.add_page_break()

    def heading(self, text: str, level: int = 1) -> None:
        self.doc.add_heading(text, level=level)

    def paragraph(self, text: str = "", *, bold_lead: str | None = None, style: str | None = None) -> None:
        p = self.doc.add_paragraph(style=style)
        if bold_lead and text.startswith(bold_lead):
            p.add_run(bold_lead).bold = True
            p.add_run(text[len(bold_lead):])
        else:
            p.add_run(text)

    def bullets(self, items: Iterable[str], *, ordered: bool = False) -> None:
        style = "List Number" if ordered else "List Bullet"
        for item in items:
            p = self.doc.add_paragraph(style=style)
            p.paragraph_format.space_after = Pt(2.5)
            p.add_run(item)

    def callout(self, title: str, text: str, *, kind: str = "info") -> None:
        fill = {"info": "EAF2FF", "warning": "FFF4E5", "risk": "FDECEC", "success": "EAF8F1"}.get(kind, "EAF2FF")
        accent = {"info": BLUE, "warning": AMBER, "risk": RED, "success": GREEN}.get(kind, BLUE)
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        cell = table.cell(0, 0)
        shade_cell(cell, fill)
        set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.name = "Aptos"
        r.font.color.rgb = RGBColor.from_string(accent)
        r.font.size = Pt(10)
        p2 = cell.add_paragraph(text)
        p2.paragraph_format.space_after = Pt(0)
        p2.runs[0].font.name = "Aptos"
        p2.runs[0].font.size = Pt(9.2)
        p2.runs[0].font.color.rgb = RGBColor.from_string(INK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)

    def table(self, caption: str, headers: Sequence[str], rows: Sequence[Sequence[object]], *, font_size: float = 7.8) -> None:
        self.table_no += 1
        cap = self.doc.add_paragraph(style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.add_run(f"Table {self.table_no}: {caption}")
        set_keep_with_next(cap)

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        header_row = table.rows[0]
        set_repeat_table_header(header_row)
        set_row_cant_split(header_row)
        for idx, header in enumerate(headers):
            shade_cell(header_row.cells[idx], NAVY)
            set_cell_text(header_row.cells[idx], header, bold=True, color=WHITE, size=font_size)
        for row_idx, values in enumerate(rows):
            added_row = table.add_row()
            set_row_cant_split(added_row)
            cells = added_row.cells
            if row_idx % 2:
                for cell in cells:
                    shade_cell(cell, LIGHT)
            for col_idx, value in enumerate(values):
                set_cell_text(cells[col_idx], "" if value is None else str(value), size=font_size)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)

    def evidence(self, *paths: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run("Evidence: " + "; ".join(paths))
        r.font.name = "Aptos"
        r.font.size = Pt(7.5)
        r.font.italic = True
        r.font.color.rgb = RGBColor.from_string(SLATE)

    def figure(self, path: Path, caption: str, *, width: float = 6.85) -> None:
        self.figure_no += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Inches(width))
        cap = self.doc.add_paragraph(style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(f"Figure {self.figure_no}: {caption}")

    def save(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(path)


def create_feature_importance_figure(metadata: dict, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig, axes = plt.subplots(1, 3, figsize=(15, 5.5), constrained_layout=True)
    palette = ["#2563EB", "#7C3AED", "#059669"]
    titles = ["Total sales", "Profit", "Stockout flag"]
    for ax, key, color, title in zip(axes, ["total_sales", "profit", "stockout_flag"], palette, titles):
        values = metadata[key]["feature_importance"][:7]
        labels = [item["feature"] for item in values][::-1]
        scores = [item["importance"] for item in values][::-1]
        ax.barh(labels, scores, color=color, edgecolor="#17365D", linewidth=0.5)
        ax.set_title(f"{title}: normalized importance", loc="left", fontsize=11, fontweight="bold")
        ax.set_xlim(0, max(scores) * 1.14 if scores else 1)
        ax.grid(axis="x", color="#D9E2F2", linewidth=0.7)
        ax.set_axisbelow(True)
        ax.spines[["top", "right", "left"]].set_visible(False)
        ax.tick_params(axis="both", labelsize=8)
        for y, score in enumerate(scores):
            ax.text(score + max(scores) * 0.015, y, f"{score:.3f}", va="center", fontsize=7.5, color="#172033")
    fig.suptitle("Saved model artifacts: most influential features", fontsize=16, fontweight="bold", x=0.02, ha="left", color="#17365D")
    fig.patch.set_facecolor("#FFFFFF")
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="#FFFFFF")
    plt.close(fig)


def create_confusion_matrix_figure(metadata: dict, path: Path) -> None:
    cm_rows = metadata["stockout_flag"]["confusion_matrix"]
    matrix = np.zeros((2, 2), dtype=int)
    for item in cm_rows:
        matrix[int(item["actual"]), int(item["predicted"])] = int(item["count"])
    fig, ax = plt.subplots(figsize=(7, 5.6), constrained_layout=True)
    image = ax.imshow(matrix, cmap="Blues", vmin=0, vmax=matrix.max())
    total = matrix.sum()
    for i in range(2):
        for j in range(2):
            value = matrix[i, j]
            ax.text(j, i, f"{value:,}\n{value / total:.1%} of test rows", ha="center", va="center", fontsize=12, fontweight="bold", color="white" if value > matrix.max() * 0.55 else "#172033")
    ax.set_xticks([0, 1], ["Predicted 0", "Predicted 1"])
    ax.set_yticks([0, 1], ["Actual 0", "Actual 1"])
    ax.set_xlabel("Prediction")
    ax.set_ylabel("Observed class")
    ax.set_title("Stockout classifier confusion matrix", loc="left", fontsize=15, fontweight="bold", color="#17365D")
    ax.spines[:].set_visible(False)
    fig.colorbar(image, ax=ax, fraction=0.045, pad=0.04, label="Test rows")
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="#FFFFFF")
    plt.close(fig)


SERVICE_ROWS = [
    ("agent.py", "Legacy/alternate", "Coordinates domain-specific dataframe tools for business analytics."),
    ("agent_loop.py", "Conditional live path", "LLM plan-act-observe loop that calls deterministic KPI, trend, quality, what-if, prediction, XAI, and root-cause tools."),
    ("ai_khata.py", "Domain analytics", "Detects and analyzes AI Khata-style transaction data."),
    ("ai_khata_tools.py", "Domain tools", "Question-answering helpers for AI Khata and transaction ledgers."),
    ("analysis_pipeline.py", "Core active", "Runs profiling, semantic mapping, planning, metrics, EDA, prediction, XAI, charts, narration, and response assembly."),
    ("analysis_planner.py", "Core active", "Chooses analysis depth and methods based on dataset evidence and query intent."),
    ("analytics_tools.py", "Legacy support", "Generic dataframe operations used by the older agent coordinator."),
    ("audit.py", "Core active", "Builds a flat audit trail from deterministic provenance receipts."),
    ("auth_service.py", "Core active", "Implements Supabase signup, email confirmation, login, refresh, JWT verification, and identity resolution."),
    ("business_leads_tools.py", "Domain tools", "Analytics helpers for business-leads datasets."),
    ("business_metrics.py", "Core active", "Calculates revenue, quantity, profit, margin, average order value, segment rankings, and query answers."),
    ("certificate.py", "Core active", "Creates SHA-256 dataset/result fingerprints and reproducibility certificates."),
    ("counterfactual.py", "Core active after modeling", "Searches for the smallest single-feature change that changes a prediction."),
    ("data_loader.py", "Support", "Loads CSV and Excel dataframes."),
    ("data_profiler.py", "Core active", "Profiles columns, missingness, uniqueness, distributions, preview rows, and semantic roles."),
    ("data_quality.py", "Core active", "Computes quality scores, EDA, IQR outliers, correlations, linear trends, and validated chart specifications."),
    ("dataset_classifier.py", "Core active", "Classifies semantic dataset type from roles and sample values."),
    ("deepanalyze_client.py", "Optional", "Calls a configured remote or local OpenAI-compatible DeepAnalyze endpoint."),
    ("domain_guard.py", "Core active when enabled", "Rejects datasets outside the retail/mart commerce domain."),
    ("health_checker.py", "Operational", "Checks dependencies and external service readiness."),
    ("intent_router.py", "Core support", "Routes natural-language questions to dataset-aware analysis intents."),
    ("llm_provider.py", "Optional", "Provider chain for OpenAI, Gemini, Anthropic, DeepAnalyze, then deterministic fallback."),
    ("llm_service.py", "Optional/secondary", "Simpler OpenAI/Gemini text-generation wrapper used by some modules."),
    ("modeling.py", "Core conditional", "Trains regression or classification candidates on an uploaded dataset and selects the best non-dummy model."),
    ("progress_bus.py", "Core active", "Publishes per-session Server-Sent Events for pipeline progress."),
    ("provenance.py", "Core active", "Builds formula, source-column, row-count, and sample-row receipts for computed values."),
    ("quality_doctor.py", "Core active", "Detects and optionally fixes duplicates, missing values, constants, and type problems."),
    ("query_planner.py", "Core active", "Converts a user question into metric, dimension, intent, filters, and report mode."),
    ("rate_limiter.py", "Core active", "In-memory sliding-window protection for authentication, upload, and message POST requests."),
    ("recommendation_engine.py", "Core support", "Generates evidence-grounded follow-up questions and recommendations."),
    ("report_composer.py", "Core active", "De-duplicates and organizes report sections, charts, tables, insights, and audit notes."),
    ("report_generator.py", "Core active", "Creates self-contained HTML and ReportLab PDF reports with tables and chart renderings."),
    ("report_narrator.py", "Core active", "Produces deterministic or provider-assisted summaries from computed facts."),
    ("response_builder.py", "Support", "Normalizes shared deterministic response payloads."),
    ("root_cause.py", "Core active on request", "Decomposes period changes by dimensions and price/volume/mix effects with receipts."),
    ("sales_tools.py", "Domain tools", "Sales-specific deterministic dataframe operations."),
    ("semantic_mapper.py", "Core active", "Maps arbitrary columns to roles and recognized dataset types, with optional LLM refinement."),
    ("session_service.py", "Core orchestrator", "Owns sessions, datasets, access checks, analysis, messages, reports, Supabase, and local fallback."),
    ("session_store.py", "Core local runtime", "Persists dataframe pickle/CSV files, metadata, and semantic maps per session and dataset."),
    ("supabase_client.py", "Core active when configured", "Custom backend REST and Storage client using the service-role credential, with local JSON/file fallback."),
    ("target_inference.py", "Core conditional", "Ranks safe regression/classification targets and rejects identifier-like candidates."),
    ("title_generator.py", "Optional", "Creates short session titles from dataset and query context."),
    ("whatif.py", "Core active on request", "Applies deterministic percentage scenarios to numeric columns and recomputes receipts."),
    ("xai.py", "Core conditional", "Uses SHAP TreeExplainer for tree models or normalized feature importance fallback, then counterfactual search."),
]


API_ROWS = [
    ("GET", "/health/live", "Public liveness check"),
    ("GET", "/api/health", "API health and configuration status"),
    ("POST", "/api/auth/signup", "Create a Supabase password user and require email confirmation"),
    ("POST", "/api/auth/login", "Exchange email/password for Supabase access and refresh tokens"),
    ("POST", "/api/auth/resend-signup", "Resend the signup confirmation link"),
    ("POST", "/api/auth/refresh", "Refresh an expired access session"),
    ("GET", "/api/auth/me", "Validate the bearer token and return the public user profile"),
    ("POST", "/api/sessions", "Create a chat/analysis session"),
    ("GET", "/api/sessions", "List recent sessions for the resolved identity"),
    ("GET", "/api/sessions/{session_id}", "Load a session with messages, datasets, agent runs, and reports"),
    ("PATCH", "/api/sessions/{session_id}", "Update session metadata such as title"),
    ("DELETE", "/api/sessions/{session_id}", "Delete an authorized session"),
    ("POST", "/api/sessions/{session_id}/datasets/upload", "Upload CSV/XLS/XLSX into a session"),
    ("GET", "/api/sessions/{session_id}/datasets", "List datasets belonging to a session"),
    ("POST", "/api/sessions/{session_id}/analyze", "Run the main analysis pipeline"),
    ("POST", "/api/sessions/{session_id}/datasets/{dataset_id}/clean", "Apply selected Data Quality Doctor fixes and re-analyze"),
    ("POST", "/api/sessions/{session_id}/datasets/{dataset_id}/verify", "Recompute and verify a reproducibility certificate"),
    ("POST", "/api/sessions/{session_id}/datasets/{dataset_id}/whatif", "Run a receipt-backed percentage scenario"),
    ("POST", "/api/sessions/{session_id}/datasets/{dataset_id}/investigate", "Run deterministic root-cause analysis"),
    ("POST", "/api/sessions/{session_id}/messages", "Ask a follow-up question against a dataset"),
    ("GET", "/api/sessions/{session_id}/progress/stream", "Stream analysis progress over Server-Sent Events"),
    ("POST", "/api/sessions/{session_id}/reports/{report_id}/renarrate", "Re-run narration without recomputing analytical facts"),
    ("GET", "/api/sessions/{session_id}/agent-runs", "List persisted agent executions"),
    ("GET", "/api/sessions/{session_id}/reports", "List reports for a session"),
    ("POST", "/api/sessions/{session_id}/reports/generate", "Generate HTML/PDF analysis outputs"),
    ("GET", "/api/reports/{report_id}/download", "Return a local file or redirect to a signed Supabase URL"),
    ("GET", "/api/datasets", "Compatibility list of recent datasets"),
    ("POST", "/api/datasets/upload", "Compatibility upload that creates a new session"),
    ("GET", "/api/datasets/{dataset_id}", "Compatibility dataset metadata lookup"),
    ("GET", "/api/datasets/{dataset_id}/profile", "Return a stored dataset profile"),
    ("POST", "/api/datasets/{dataset_id}/ask", "Compatibility direct question endpoint"),
    ("DELETE", "/api/datasets/{dataset_id}", "Compatibility dataset deletion"),
    ("GET", "/api/storage/status", "Report Supabase or local persistence mode"),
    ("POST", "/api/upload", "Legacy upload/profile endpoint"),
    ("GET", "/api/session/{session_id}", "Legacy session lookup"),
    ("DELETE", "/api/session/{session_id}", "Legacy session deletion"),
    ("POST", "/api/analyze/upload", "Legacy immediate full analysis"),
    ("POST", "/api/analyze/query", "Legacy query against an uploaded session"),
]


DATASET_COLUMNS = [
    ("store_id", "int64", "Store identifier; excluded from saved models because names ending in _id are treated as identifiers."),
    ("region", "int64", "Encoded region category. The repository does not include a code-to-label mapping."),
    ("city", "int64", "Encoded city category; seven observed codes."),
    ("category", "int64", "Encoded product category; four observed codes."),
    ("subcategory", "int64", "Encoded product subcategory; ten observed codes."),
    ("unit_price", "float64", "Unit selling price."),
    ("quantity", "int64", "Units in the transaction."),
    ("discount", "float64", "Discount fraction applied to the transaction."),
    ("total_sales", "float64", "Post-discount sales value. Approximately price_qty - discount_value, rounded to cents."),
    ("profit", "float64", "Transaction profit target used by one saved regression artifact."),
    ("customer_type", "int64", "Encoded customer segment."),
    ("payment_method", "int64", "Encoded payment method."),
    ("online_order", "int64", "Binary indicator for online order status."),
    ("stockout_flag", "int64", "Binary stockout target; 7,053 positive and 27,947 negative rows."),
    ("weekday", "int64", "Encoded weekday."),
    ("month", "int64", "Month number."),
    ("year", "int64", "Calendar year."),
    ("hour", "int64", "Hour of day."),
    ("price_qty", "float64", "Engineered value equal to unit_price multiplied by quantity."),
    ("discount_value", "float64", "Engineered value equal to price_qty multiplied by discount."),
    ("profit_margin", "float64", "Engineered profit ratio."),
]


FRONTEND_ROWS = [
    ("app/page.tsx", "/", "Marketing landing page"),
    ("app/about/page.tsx", "/about", "Project principles, pipeline, and technology overview"),
    ("app/features/page.tsx", "/features", "Feature catalog"),
    ("app/signup/page.tsx", "/signup", "Validated signup and email-confirmation waiting state"),
    ("app/login/page.tsx", "/login", "Login and confirmed-email feedback"),
    ("app/dashboard/page.tsx", "/dashboard", "Protected shell that renders the analysis workspace"),
    ("components/dashboard/DashboardApp.tsx", "Dashboard", "Upload, session, chat, analysis, reports, XAI, and view-state orchestration"),
    ("components/DropZone.tsx", "Upload", "Drag/drop and extension checks for CSV/XLS/XLSX"),
    ("components/ThinkingTrace.tsx", "Progress", "Renders live Server-Sent Events pipeline steps"),
    ("components/dashboard/KpiCard.tsx", "Evidence", "KPI value plus provenance drill-down"),
    ("components/dashboard/VerificationPanel.tsx", "Evidence", "Audit receipt inspection"),
    ("components/dashboard/CertificateCard.tsx", "Evidence", "Reproducibility certificate validation"),
    ("components/dashboard/QualityDoctorPanel.tsx", "Quality", "Issue selection and deterministic cleaning"),
    ("components/dashboard/WhatIfPanel.tsx", "Analysis", "Interactive percentage scenario inputs"),
    ("components/dashboard/RootCausePanel.tsx", "Analysis", "Driver decomposition and price/volume/mix display"),
    ("components/dashboard/AgentTracePanel.tsx", "Agent", "Plan-act-observe tool trace"),
]


DEPENDENCY_ROWS = [
    ("FastAPI", "0.135.2", "Direct", "HTTP API, dependency injection, request validation, routing"),
    ("Uvicorn", "0.42.0", "Runtime", "ASGI development/production process"),
    ("Pandas", "2.3.3", "Direct, 31 backend files", "Dataframes, aggregation, profiling, CSV/Excel operations"),
    ("NumPy", "2.3.5", "Direct", "Numeric arrays, polyfit trends, model/XAI support"),
    ("scikit-learn", "1.8.0", "Direct", "Preprocessing, candidate models, metrics, train/test split"),
    ("SHAP", "0.51.0", "Conditional direct", "TreeExplainer local explanations"),
    ("SciPy", "1.16.3", "Declared/supporting", "Scientific stack dependency; not directly imported by active app files"),
    ("OpenPyXL", "3.1.5", "Indirect", "Pandas Excel reader for XLSX"),
    ("ReportLab", "4.5.1 installed", "Direct", "PDF generation and chart drawings"),
    ("Jinja2", "3.1.6", "Direct", "HTML report templating"),
    ("HTTPX", "0.28.1", "Direct", "Supabase Auth/REST/Storage, Gemini HTTP, DeepAnalyze, health checks"),
    ("Pydantic", "2.12.5", "Direct", "API and semantic/query schemas"),
    ("pydantic-settings", "2.13.1", "Direct", "Environment configuration"),
    ("python-multipart", "0.0.26", "Framework support", "Multipart file uploads"),
    ("PyJWT", "2.13.0", "Direct", "Legacy HS256 Supabase JWT validation"),
    ("OpenAI", "1.109.1", "Optional direct", "Report/query narration and agent planning"),
    ("google-generativeai", "0.8.5", "Optional direct", "Gemini provider"),
    ("Anthropic", "0.100.0", "Optional direct", "Claude provider"),
    ("Supabase Python", "Not installed", "Not required by current code", "The application uses a custom HTTPX Supabase client instead"),
    ("pytest", "9.0.2", "Test", "Backend test runner"),
    ("pytest-asyncio", "1.4.0 installed", "Test", "Async endpoint/service tests"),
    ("pypdf", "6.14.2", "Test/report validation", "PDF inspection"),
    ("Next.js", "15.5.20 installed", "Direct frontend", "App Router, production build, local HTTPS dev server"),
    ("React / React DOM", "19.2.7 installed", "Direct frontend", "UI rendering and client state"),
    ("TypeScript", "5.9.3", "Build", "Static frontend types"),
    ("Tailwind CSS", "4.1.11", "Direct frontend", "Utility styling"),
    ("Motion", "12.42.2 installed", "Direct frontend", "Animations"),
    ("Lucide React", "0.553.0", "Direct frontend", "Icons"),
    ("react-markdown", "10.1.0", "Direct frontend", "Assistant/report markdown rendering"),
    ("ESLint", "9.39.1", "Quality", "Frontend static analysis"),
]


ENV_ROWS = [
    ("Application", "ENVIRONMENT, APP_VERSION, ENABLE_OPENAPI_DOCS, REQUEST_TIMEOUT_SECONDS", "Runtime behavior and API documentation"),
    ("Transport", "CORS_ORIGINS, CORS_ORIGIN_REGEX, TRUSTED_HOSTS, SECURE_HEADERS_ENABLED, HTTPS_REDIRECT", "Browser/API transport controls"),
    ("Rate limits", "RATE_LIMIT_ENABLED, RATE_LIMIT_REQUESTS, RATE_LIMIT_WINDOW_SECONDS, RATE_LIMIT_PER_MINUTE", "Abuse protection thresholds"),
    ("Supabase", "SUPABASE_URL, SUPABASE_ANON_KEY, SUPABASE_SERVICE_ROLE_KEY, SUPABASE_JWT_SECRET", "Authentication, PostgreSQL REST, Storage, JWT validation"),
    ("Supabase storage", "SUPABASE_DATASET_BUCKET, SUPABASE_REPORT_BUCKET, SUPABASE_HEALTH_TIMEOUT_SECONDS", "Private bucket names and health checks"),
    ("URLs", "BACKEND_BASE_URL, FRONTEND_BASE_URL, NEXT_PUBLIC_DATAVERSE_API_URL, NEXT_PUBLIC_API_URL", "API routing and confirmation redirects"),
    ("Uploads/modeling", "MAX_UPLOAD_SIZE_MB, RETAIL_ONLY_UPLOADS, AUTO_TRAIN_TARGET_CONFIDENCE, MIN_ROWS_FOR_PREDICTION", "Dataset admission and modeling thresholds"),
    ("LLM control", "USE_LLM_NARRATION, LLM_PROVIDER, REPORT_NARRATOR_TIMEOUT_SECONDS, INTENT_LLM_PROVIDER, INTENT_LLM_TIMEOUT", "Optional provider use"),
    ("OpenAI", "OPENAI_API_KEY, OPENAI_API_BASE, OPENAI_CHAT_MODEL, OPENAI_INTENT_MODEL", "OpenAI narration/planning"),
    ("Gemini", "GEMINI_API_KEY, GEMINI_API_BASE, GEMINI_MODEL, GEMINI_REPORT_MODEL, GOOGLE_API_KEY", "Google provider configuration"),
    ("Anthropic", "ANTHROPIC_API_KEY, CLAUDE_MODEL", "Claude provider configuration"),
    ("DeepSeek", "DEEPSEEK_API_KEY, DEEPSEEK_API_BASE, DEEPSEEK_INTENT_MODEL", "Intent parsing provider"),
    ("DeepAnalyze", "DEEPANALYZE_API_KEY, DEEPANALYZE_API_BASE, DEEPANALYZE_LOCAL_BASE_URL, DEEPANALYZE_MODEL", "Remote/local OpenAI-compatible provider"),
    ("Logging", "LOG_DIR, LOG_LEVEL, LOG_JSON", "Operational logs"),
    ("Optional infrastructure", "DATABASE_URL, REDIS_URL, CELERY_BROKER_URL, CELERY_RESULT_BACKEND, STORAGE_TYPE", "Configured in settings but not part of the active session flow"),
    ("Optional object storage", "MINIO_*, AWS_*", "Configured in settings but inactive in the current Supabase/local path"),
    ("Observability", "SENTRY_DSN, SENTRY_TRACES_SAMPLE_RATE, SENTRY_PROFILES_SAMPLE_RATE", "Declared configuration; no active Sentry initialization found"),
]


def add_cover(builder: ReportBuilder) -> None:
    doc = builder.doc
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DATAVERSE AI")
    run.font.name = "Aptos Display"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Complete Current-Folder Technical Audit")
    run.font.name = "Aptos Display"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Every project file, model, dataset, library, architecture path, API, Supabase component, test, artifact, and deployment input")
    r.font.name = "Aptos"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(SLATE)

    doc.add_paragraph().paragraph_format.space_after = Pt(40)
    info = doc.add_table(rows=5, cols=2)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Repository", str(ROOT)),
        ("Audit date", "14 July 2026"),
        ("Audience", "Technical reviewers, project supervisors, developers, and defense panel"),
        ("Evidence basis", "Whole-folder file inventory, source code, manifests, models, datasets, generated artifacts, tests, lint, and production build"),
        ("Document status", "Current-folder audited implementation and file-reference report"),
    ]
    for idx, (label, value) in enumerate(rows):
        shade_cell(info.cell(idx, 0), NAVY)
        set_cell_text(info.cell(idx, 0), label, bold=True, color=WHITE, size=9)
        set_cell_text(info.cell(idx, 1), value, size=9)

    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("This report describes the complete current folder. Project files are itemized; third-party, cache, private session, and certificate directories are counted and safely summarized.")
    r.font.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(SLATE)
    builder.page_break()


def build_report() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    SUPPORT_DIR.mkdir(parents=True, exist_ok=True)
    metadata = json.loads(MODEL_METADATA_PATH.read_text(encoding="utf-8"))
    feature_path = SUPPORT_DIR / "saved_model_feature_importance.png"
    confusion_path = SUPPORT_DIR / "stockout_confusion_matrix.png"
    create_feature_importance_figure(metadata, feature_path)
    create_confusion_matrix_figure(metadata, confusion_path)
    inventory = build_inventory(ROOT, SUPPORT_DIR)
    inventory_summary = inventory["summary"]
    inventory_files = inventory["files"]
    excluded_directories = inventory["excluded_directories"]

    b = ReportBuilder()
    add_cover(b)

    b.heading("Document control", 1)
    b.table(
        "Document control and interpretation rules",
        ["Item", "Value"],
        [
            ("Scope", "The complete current FINAL3 folder, including project-owned files, generated artifacts, and aggregate accounting for dependency/runtime directories"),
            ("Primary branch state", "Working tree as inspected on 14 July 2026"),
            ("Active", "Imported or reached by the current Next.js to FastAPI session flow"),
            ("Conditional", "Runs only when a feature is requested or required configuration exists"),
            ("Optional", "Present in code/configuration but not required for deterministic analysis"),
            ("Legacy/compatibility", "Still mounted or tested, but not the main frontend path"),
            ("Not established", "No repository evidence proves external production configuration or live service state"),
        ],
        font_size=8.5,
    )
    b.heading("Table of contents", 1)
    toc = b.doc.add_paragraph()
    add_toc(toc)
    b.page_break()

    b.heading("1. Technical summary", 1)
    b.paragraph(
        "DataVerse AI is a full-stack, deterministic-first dataset analysis system. A Next.js 15 frontend communicates with a FastAPI backend. The backend accepts CSV and Excel files, profiles and semantically maps their columns, computes business metrics with Pandas and NumPy, trains scikit-learn models when requested and safe, explains eligible models with SHAP or feature importance, generates HTML/PDF reports, and persists identity/data through Supabase or a local fallback."
    )
    b.table(
        "Most important implementation findings",
        ["Finding", "Confirmed implementation", "Implication"],
        [
            ("Deterministic calculations", "Pandas, NumPy, and scikit-learn compute numbers. LLMs only plan/refine language.", "The numerical path works without any LLM key."),
            ("Two model paths", "Saved retail artifacts exist, while the live application trains models per uploaded dataset.", "Saved .pkl files are evidence artifacts, not the current inference endpoint."),
            ("Saved model quality", "Ridge is saved for total_sales and profit; RandomForestClassifier is saved for stockout_flag.", "Metrics must be interpreted with leakage and class-imbalance caveats."),
            ("Real authentication", "Supabase Auth handles password storage, confirmation links, login, refresh, and user lookup.", "Application code never stores plaintext passwords."),
            ("Persistence", "Supabase stores metadata and private files when configured; local JSON/filesystem fallback otherwise.", "A working local app does not prove Supabase is configured."),
            ("Verified features", "Receipts, SHA-256 certificates, quality doctor, what-if, root-cause, XAI, and reports are implemented.", "Outputs can include traceable formulas and reproducibility checks."),
            ("Current verification", "214 backend tests, frontend lint, and the production Next.js build pass.", "The inspected codebase is internally consistent at the tested level."),
        ],
        font_size=8.1,
    )
    b.callout(
        "Critical interpretation",
        "The saved total_sales model reports R2 = 1.0 because engineered inputs price_qty and discount_value almost exactly reconstruct total_sales. This is target leakage for forecasting future sales and should not be presented as real predictive performance.",
        kind="risk",
    )
    b.evidence("models/model_metadata.json", "scripts/train_retail_mart.py", "dataverse_backend/app/services/modeling.py", "test/build outputs from this audit")

    b.heading("2. Scope, audit method, and evidence", 1)
    b.paragraph("The report answers: what is present anywhere in the current FINAL3 folder, what each project file is for, which components participate in the running application, and which items are historical, generated, optional, external, cached, private, or not part of production execution?")
    b.bullets([
        "Inspected source code under frontend/, dataverse_backend/app/, scripts/, models/, supabase/, and tests/.",
        "Read package.json, package-lock.json, all backend requirement manifests, Dockerfiles, environment examples, and SQL migrations.",
        "Profiled data/retail_mart_processed_v1.csv directly with Pandas.",
        "Loaded the three local serialized scikit-learn pipelines to confirm estimator types and pipeline steps.",
        "Compared saved metrics and feature importance with the training script and calculated leakage checks.",
        "Enumerated FastAPI decorators and frontend API calls.",
        "Ran the full backend test suite, frontend lint, and Next.js production build.",
        "Generated a file-by-file inventory for all project-owned and generated files that can be safely reported.",
        "Counted third-party dependencies, Git internals, caches, browser traces, private session storage, temporary files, and local certificate material without reproducing their contents.",
    ])
    b.callout("Evidence boundary", "No secret values were printed or embedded. The audit does not prove the state of the external Supabase dashboard, email provider, DNS, hosted deployment, or production environment variables.", kind="warning")

    b.heading("2.1 Whole-folder coverage", 2)
    b.table(
        "Current FINAL3 folder accounting",
        ["Measure", "Count or size", "Interpretation"],
        [
            ("Safely itemized files", f"{inventory_summary['audited_files']:,}", "Every item appears in Appendix D and in the companion CSV/JSON inventory."),
            ("Itemized file size", format_bytes(inventory_summary["audited_bytes"]), "Source, documentation, data/model artifacts, images, reports, logs, and configuration files."),
            ("Readable text lines", f"{inventory_summary['audited_text_lines']:,}", "Counted from non-sensitive text files; environment secrets were not opened for inventory content."),
            ("Git-tracked itemized files", f"{inventory_summary['tracked_files']:,}", "Files currently recorded in Git."),
            ("Local/generated/untracked itemized files", f"{inventory_summary['local_generated_untracked_files']:,}", "Working outputs, screenshots, reports, logs, or uncommitted project files."),
            ("Safely summarized excluded files", f"{inventory_summary['excluded_files']:,}", "Dependency internals, caches, Git database, private runtime state, temporary conversions, and certificate material."),
            ("Physical files accounted for", f"{inventory_summary['physical_files_accounted_for']:,}", "Itemized files plus every file counted inside intentionally summarized directories."),
        ],
        font_size=7.7,
    )
    category_rows = sorted(
        inventory_summary["categories"].items(),
        key=lambda item: (CATEGORY_ORDER.get(item[0], 99), item[0]),
    )
    b.table(
        "Itemized files by repository responsibility",
        ["Repository responsibility", "Files", "Coverage"],
        [(category, f"{count:,}", "File-by-file inventory in Appendix D") for category, count in category_rows],
        font_size=7.5,
    )
    b.callout(
        "How to read the complete inventory",
        "The Word/PDF appendix explains every safely reportable file. The companion CSV is the fastest searchable reference. Installed package internals are not copied into the report because package.json, lockfiles, requirement manifests, and dependency tables are the correct auditable description of those third-party libraries.",
        kind="info",
    )

    b.heading("3. System architecture and runtime flow", 1)
    b.paragraph("The active architecture has six essential parts: the user, Next.js frontend, FastAPI API, two analysis agents, deterministic analytics engine, and Supabase. Results return through the same authenticated API path used for requests.")
    if ARCHITECTURE_IMAGE.exists():
        b.figure(ARCHITECTURE_IMAGE, "DataVerse AI system architecture based on used components", width=7.0)
        b.evidence("output/architecture/dataverse-production-architecture.png", "frontend/lib/dataverse-api.ts", "dataverse_backend/app/main.py", "dataverse_backend/app/services/session_service.py")

    b.heading("3.1 End-to-end user flow", 2)
    b.table(
        "Runtime processing stages",
        ["Stage", "Component", "Processing performed", "Main output"],
        [
            ("1", "Supabase Auth via FastAPI", "Signup validates password, sends confirmation link, logs in confirmed users, refreshes sessions.", "JWT access token, refresh token, public user profile"),
            ("2", "Next.js dashboard", "Creates a session and uploads CSV/XLS/XLSX over multipart HTTP.", "Session ID and dataset upload request"),
            ("3", "FastAPI + SessionService", "Validates size/extension, checks ownership, parses data, stores file and metadata.", "Dataset ID, profile, semantic map"),
            ("4", "DatasetAgent", "Normalizes headers, profiles shape/types/missingness, and computes initial quality evidence.", "Profile and data-quality payload"),
            ("5", "AnalystAgent + AnalysisPipeline", "Plans intent, computes metrics/EDA/trends/correlations/outliers, and conditionally trains a model/XAI.", "Facts, tables, charts, receipts, model evidence"),
            ("6", "ReportNarrator/Generator", "Organizes deterministic facts and optionally improves wording with an LLM.", "Chat answer, HTML report, PDF report"),
            ("7", "Supabase/local persistence", "Stores sessions/messages/metadata and private dataset/report files.", "Durable metadata and signed/local report access"),
            ("8", "Frontend evidence panels", "Displays KPIs, charts, quality issues, agent trace, root cause, what-if, certificates, and reports.", "Interactive analysis workspace"),
        ],
        font_size=7.5,
    )

    b.heading("3.2 Active and compatibility API paths", 2)
    b.paragraph("The main frontend uses /api/sessions/* and /api/datasets/* routes. Legacy /api/upload and /api/analyze/* routes remain mounted and are covered by tests. This increases compatibility but also increases the maintenance and security-review surface.")

    b.page_break()
    b.heading("4. Repository and technology stack", 1)
    b.table(
        "Repository areas",
        ["Area", "Purpose", "Status"],
        [
            ("frontend/", "Next.js App Router pages, dashboard components, API client, auth context, local HTTPS launcher", "Active"),
            ("dataverse_backend/app/api/", "FastAPI route definitions and request schemas", "Active plus compatibility routes"),
            ("dataverse_backend/app/agents/", "DatasetAgent and AnalystAgent wrappers; older agents also remain", "Mixed active/legacy"),
            ("dataverse_backend/app/services/", "Core analytics, orchestration, persistence, security, reporting, and AI logic", "Primary implementation"),
            ("dataverse_backend/tests/", "Backend unit, integration, security, API, reporting, and agentic tests", "214 passing"),
            ("supabase/migrations/", "Chat/session database and storage schema", "Deployment input"),
            ("data/", "Processed 35,000-row retail dataset", "Training/evaluation input"),
            ("models/", "Three saved sklearn pipelines and metadata", "Offline artifacts; not loaded by live app"),
            ("scripts/", "Launch, smoke test, model training, academic documents, and this technical report", "Tooling"),
            ("docs/", "Architecture, setup, evaluation, deployment, and historical completion documents", "Documentation; some statements may be stale"),
        ],
        font_size=8,
    )

    b.heading("4.1 Primary stack", 2)
    b.table(
        "Primary application technologies",
        ["Layer", "Technology", "Installed/deployed role"],
        [
            ("Frontend", "Next.js 15.5.20, React 19.2.7, TypeScript 5.9.3", "App Router pages and client dashboard"),
            ("Styling", "Tailwind CSS 4.1.11, clsx, tailwind-merge, Motion, Lucide", "Responsive UI, animation, icons"),
            ("Backend", "Python 3.12, FastAPI 0.135.2, Uvicorn 0.42.0", "REST API and ASGI server"),
            ("Data", "Pandas 2.3.3, NumPy 2.3.5, OpenPyXL 3.1.5", "Dataframes, numerical processing, Excel input"),
            ("Machine learning", "scikit-learn 1.8.0, SHAP 0.51.0", "Preprocessing, candidates, metrics, explanations"),
            ("Reports", "Jinja2 3.1.6, ReportLab 4.5.1", "HTML/PDF output"),
            ("Identity/data platform", "Supabase Auth, PostgreSQL REST, Storage", "Accounts, metadata, private file objects"),
            ("External text AI", "OpenAI, Gemini, Anthropic, DeepAnalyze, DeepSeek configuration", "Optional planning and narration only"),
            ("Verification", "pytest, ESLint, Next build", "Backend and frontend quality gates"),
        ],
        font_size=8.1,
    )

    b.page_break()
    b.heading("5. Frontend implementation", 1)
    b.paragraph("The frontend is a Next.js App Router application. Public marketing/authentication pages are statically generated. The dashboard is a client-heavy workspace that orchestrates sessions, uploads, chat, evidence panels, and report links through a typed API client.")
    b.table("Frontend pages and important components", ["File", "Surface", "Responsibility"], FRONTEND_ROWS, font_size=7.7)
    b.heading("5.1 Authentication state", 2)
    b.bullets([
        "AuthProvider calls only the FastAPI authentication routes; the browser does not call Supabase directly.",
        "The access token, refresh token, and display session are kept in browser localStorage under dataverse.token, dataverse.refreshToken, and dataverse.session.",
        "On reload, /api/auth/me validates the access token. A failed access token triggers /api/auth/refresh when a refresh token exists.",
        "The dashboard route redirects unauthenticated users to /login at the client layer.",
        "API requests send Authorization: Bearer <token>. A legacy X-Dataverse-User workspace ID is also generated and sent.",
    ])
    b.heading("5.2 Local HTTPS", 2)
    b.paragraph("npm run dev starts FastAPI on 127.0.0.1:8000 and Next.js with --experimental-https on 127.0.0.1:3000. Browser API traffic uses the secure same-origin /backend/api proxy, while Next.js communicates server-side with the local HTTP backend.")
    b.evidence("frontend/scripts/dev.mjs", "frontend/next.config.ts", "frontend/lib/apiConfig.ts", "frontend/lib/auth.tsx", "frontend/lib/dataverse-api.ts")

    b.heading("6. Backend API and orchestration", 1)
    b.paragraph("FastAPI mounts authentication, core, session, dataset, report, storage, and legacy analysis routers. SessionService is the central application service: it enforces session ownership, loads datasets, persists messages and agent runs, invokes AnalysisPipeline, and generates reports.")
    b.table(
        "Backend boundary controls",
        ["Control", "Implementation"],
        [
            ("CORS", "Configured allow-list with optional deployment regex; credentials allowed"),
            ("Compression", "GZip for responses of 500 bytes or more"),
            ("Rate limiting", "POST auth/upload/message routes use an in-memory sliding window"),
            ("Security headers", "nosniff, frame denial, strict-origin referrer policy, disabled camera/microphone/geolocation"),
            ("Ownership", "SessionService.ensure_access blocks user ID mismatches"),
            ("Errors", "Development returns diagnostic details; production returns a generic unexpected-error message"),
            ("Progress", "Per-session Server-Sent Events stream reports pipeline stages"),
        ],
        font_size=8.2,
    )
    b.heading("6.1 Complete API catalog", 2)
    b.table("FastAPI endpoint catalog", ["Method", "Path", "Purpose"], API_ROWS, font_size=7.1)
    b.evidence("dataverse_backend/app/main.py", "dataverse_backend/app/api/*.py")

    b.page_break()
    b.heading("7. Agents and analytical pipeline", 1)
    b.heading("7.1 DatasetAgent", 2)
    b.paragraph("DatasetAgent owns ingestion. It rejects empty files, files larger than MAX_UPLOAD_SIZE_MB (50 MB by default), and extensions other than .csv, .xlsx, or .xls. It parses the file, normalizes blank/duplicate headers, stores the dataframe, and returns profile and quality evidence.")
    b.heading("7.2 AnalystAgent", 2)
    b.paragraph("AnalystAgent wraps AnalysisPipeline. It receives the dataframe, semantic map, query, prediction target/task preferences, XAI flag, and report flag. It returns deterministic analysis facts plus optional narrative/report outputs.")
    b.heading("7.3 AnalysisPipeline stage inventory", 2)
    b.table(
        "AnalysisPipeline methods and algorithms",
        ["Stage", "Primary module", "Method"],
        [
            ("Coercion/profile", "data_profiler.py", "Type detection, missing masks, preview, distributions, cardinality, semantic roles"),
            ("Dataset type", "semantic_mapper.py + dataset_classifier.py", "Rule-based role and sample inspection; optional LLM refinement"),
            ("Planning", "query_planner.py + analysis_planner.py", "Intent, metric, dimension, filters, report mode, prediction need"),
            ("Business metrics", "business_metrics.py", "SUM, derived profit/margin/AOV, date/dimension groupings, rankings"),
            ("Quality", "data_quality.py + quality_doctor.py", "Missingness, duplicates, constants, high cardinality, type issues, fixes"),
            ("EDA", "data_quality.py", "describe(), frequency tables, histograms, IQR outliers"),
            ("Relationships", "data_quality.py", "Pearson correlation matrix for numeric columns"),
            ("Trends", "data_quality.py", "Date aggregation and NumPy first-degree polyfit slope"),
            ("Prediction", "target_inference.py + modeling.py", "Safe target ranking, preprocessing, holdout comparison of candidate estimators"),
            ("Explainability", "xai.py + counterfactual.py", "Tree SHAP or importance fallback, plus single-feature counterfactual search"),
            ("Verification", "provenance.py + audit.py + certificate.py", "Formula receipts and SHA-256 fingerprints"),
            ("Advanced analysis", "root_cause.py + whatif.py", "Period driver decomposition and deterministic scenario recomputation"),
            ("Reporting", "report_composer.py + report_generator.py", "De-duplicated HTML/PDF report with validated charts/tables"),
        ],
        font_size=7.5,
    )

    b.heading("7.4 Dataset types recognized", 2)
    b.paragraph("Semantic classification supports mart_sales, retail_sales, invoice_sales, ecommerce_orders, pos_transactions, transaction_ledger, inventory, food_dataset, customer_sales, and generic_tabular. RETAIL_ONLY_UPLOADS defaults to true, so production-like configuration can reject unsupported domains even though generic analysis code exists.")
    b.evidence("dataverse_backend/app/agents/dataset_agent.py", "dataverse_backend/app/agents/analyst_agent.py", "dataverse_backend/app/services/analysis_pipeline.py", "dataverse_backend/app/services/semantic_mapper.py")

    b.heading("8. Dataset inputs and processing", 1)
    b.heading("8.1 User-uploaded datasets", 2)
    b.table(
        "Supported upload behavior",
        ["Property", "Implementation"],
        [
            ("Formats", "CSV, XLSX, XLS"),
            ("Default maximum", "50 MB"),
            ("CSV parsing", "Delimiter sniffing, sectioned-table detection, repair fallback for inconsistent rows"),
            ("Excel parsing", "Pandas read_excel through the installed Excel engine"),
            ("Header handling", "Whitespace cleanup and generated names for blank columns"),
            ("Persistence", "Per-dataset pickle and CSV, metadata JSON, semantic-map JSON; Supabase private object copy when configured"),
            ("Validation", "Non-empty dataframe, supported extension, domain guard when enabled"),
            ("Model minimum", "At least 30 rows for runtime prediction by default"),
        ],
        font_size=8.2,
    )

    b.heading("8.2 Included retail training dataset", 2)
    b.table(
        "retail_mart_processed_v1.csv profile",
        ["Measure", "Observed value"],
        [
            ("Rows", "35,000"),
            ("Columns", "21"),
            ("File size", "3,113,988 bytes (about 2.97 MiB)"),
            ("In-memory dataframe size", "5,880,132 bytes (about 5.61 MiB)"),
            ("Missing cells", "0"),
            ("Duplicate rows", "0"),
            ("total_sales", "Mean 64.2475; median 46.32; min 0; max 495.44; 14,187 unique values"),
            ("profit", "Mean 9.6268; median 6.34; min 0; max 101.63; 4,186 unique values"),
            ("stockout_flag", "7,053 positive (20.15%) and 27,947 negative (79.85%)"),
            ("Data encoding", "All 21 columns are numeric; category label mappings are not included"),
        ],
        font_size=8.4,
    )
    b.table("Included retail dataset column dictionary", ["Column", "dtype", "Meaning and caveat"], DATASET_COLUMNS, font_size=7.4)
    b.callout("Dataset limitation", "The repository contains only the processed numeric dataset. Without the raw source and encoding dictionary, region, city, category, subcategory, customer_type, and payment_method codes cannot be translated back to business labels.", kind="warning")
    b.evidence("data/retail_mart_processed_v1.csv", "scripts/train_retail_mart.py")

    b.page_break()
    b.heading("9. Machine-learning models", 1)
    b.heading("9.1 Live per-upload modeling", 2)
    b.paragraph("The live application does not use a fixed prediction model. When prediction is explicitly requested, implied by the query, or enabled with a target confidence at or above 0.65, modeling.py builds a model from the uploaded dataset.")
    b.table(
        "Live model training design",
        ["Aspect", "Regression", "Classification"],
        [
            ("Target detection", "Numeric, sufficiently high-cardinality target", "Boolean/low-cardinality numeric or categorical target"),
            ("Dummy baseline", "DummyRegressor(strategy=mean)", "DummyClassifier(strategy=most_frequent)"),
            ("Candidate 1", "Ridge(random_state=42)", "LogisticRegression(max_iter=1000, class_weight=balanced)"),
            ("Candidate 2", "RandomForestRegressor(80 trees, max_depth=8, random_state=42)", "RandomForestClassifier(80 trees, max_depth=8, class_weight=balanced, random_state=42)"),
            ("Split", "75% train / 25% test, random_state=42", "75% train / 25% test, stratified when each class has at least two rows"),
            ("Selection", "Highest test R2 among non-dummy candidates", "Highest weighted test F1 among non-dummy candidates"),
            ("Metrics", "R2, RMSE, MAE", "Accuracy, weighted F1, confusion matrix"),
            ("Numeric preprocessing", "Median imputation + StandardScaler", "Median imputation + StandardScaler"),
            ("Categorical preprocessing", "Most-frequent imputation + one-hot encoding", "Most-frequent imputation + one-hot encoding"),
            ("Safety gates", "Minimum 30 rows, identifiers/constants/high-cardinality and name-overlap leakage excluded, max 200 estimated one-hot features", "Same gates plus imbalance warning when one class is at least 90%"),
        ],
        font_size=7.4,
    )

    b.heading("9.2 Saved retail model artifacts", 2)
    b.paragraph("scripts/train_retail_mart.py trained and serialized three sklearn Pipeline objects. Each pipeline contains preprocess and model steps. A repository search found no application import or load of these .pkl files, so they are separate training artifacts rather than the live prediction implementation.")
    b.table(
        "Saved model summary",
        ["Target", "Task", "Selected estimator", "Test metrics", "Artifact"],
        [
            ("total_sales", "Regression", "Ridge", "R2 1.000000; RMSE 0.010614; MAE 0.007332", "models/total_sales_model.pkl"),
            ("profit", "Regression", "Ridge", "R2 0.753999; RMSE 5.019354; MAE 3.190698", "models/profit_model.pkl"),
            ("stockout_flag", "Classification", "RandomForestClassifier", "Accuracy 0.638286; weighted F1 0.655691; weighted precision 0.677070; weighted recall 0.638286", "models/stockout_flag_model.pkl"),
        ],
        font_size=7.8,
    )
    b.figure(feature_path, "Top normalized feature importance recorded for each saved model", width=7.0)
    b.paragraph("The feature-importance panels show that total_sales is dominated by price_qty (0.875938) and discount_value (0.123233). Those two engineered features reconstruct total_sales within about 0.005 due to rounding. Profit is dominated by total_sales (0.471716) and price_qty (0.420806). The stockout classifier spreads importance across profit_margin, total_sales, unit_price, profit, price_qty, discount_value, and time/category fields.")

    b.heading("9.3 Saved model feature sets", 2)
    for target in ("total_sales", "profit", "stockout_flag"):
        entry = metadata[target]
        b.heading(target, 3)
        b.paragraph(", ".join(entry["features"]))

    b.heading("9.4 Stockout-classification behavior", 2)
    b.figure(confusion_path, "Stockout RandomForestClassifier confusion matrix", width=5.8)
    b.paragraph("Although weighted metrics look moderate, the model identifies only 464 of 1,763 actual stockouts. Positive-class recall is 26.3% and positive-class precision is 19.9%. This model is not suitable for operational stockout decisions without better features, threshold tuning, calibration, and external validation.")
    b.callout("Model validation conclusion", "The saved artifacts demonstrate a working training pipeline, not production-ready forecasting. total_sales is leaked by engineered inputs, profit has moderate holdout fit, and stockout detection performs poorly on the minority class.", kind="risk")
    b.evidence("scripts/train_retail_mart.py", "models/model_metadata.json", "models/*.pkl", "dataverse_backend/app/services/modeling.py", "dataverse_backend/app/services/target_inference.py")

    b.page_break()
    b.heading("10. Explainable AI and decision tools", 1)
    b.table(
        "Explainability and verification mechanisms",
        ["Capability", "Method", "Output", "Limitation"],
        [
            ("Global feature importance", "Normalized tree importance or absolute model coefficients", "Ranked feature list", "Association with model output, not causality"),
            ("Local SHAP", "SHAP TreeExplainer on up to 10 transformed test rows for tree-compatible models", "Top signed contributors per sample", "Not attempted for Ridge/Logistic models"),
            ("Fallback XAI", "Saved feature importance when SHAP is missing or fails", "Top features and plain-language explanation", "No local attribution"),
            ("Counterfactual XAI", "Deterministic single-feature search", "Smallest found change that changes the prediction", "Does not prove feasibility or causal effect"),
            ("What-if", "Apply percentage change to a numeric column and recompute business metrics", "Baseline, scenario, delta, receipts", "Scenario is mechanical, not a forecast"),
            ("Root cause", "Period comparison, dimension contribution ranking, price/volume/mix split", "Drivers, chart, receipts, narrative", "Descriptive decomposition, not causal identification"),
            ("Data Quality Doctor", "Rule-based duplicate/missing/constant/type checks", "Issues, proposed fixes, before/after stats", "Automated cleaning may change analytical meaning"),
            ("Reproducibility certificate", "SHA-256 fingerprints of data and selected results", "Tamper-evident certificate and verification response", "Proves repeatability under the same code/data, not real-world correctness"),
            ("Provenance receipts", "Formula, operation, source columns, row count, samples", "Audit trail for KPIs and analytical facts", "Sample rows are evidence, not a full lineage platform"),
        ],
        font_size=7.35,
    )
    b.evidence("dataverse_backend/app/services/xai.py", "counterfactual.py", "whatif.py", "root_cause.py", "quality_doctor.py", "certificate.py", "provenance.py", "audit.py")

    b.page_break()
    b.heading("11. LLM and language-model integration", 1)
    b.paragraph("LLMs are optional. The application is designed so provider failure returns to deterministic templates rather than blocking calculations. Numeric facts are computed before narration and prompts instruct providers not to invent values.")
    b.table(
        "Configured language-model providers",
        ["Provider", "Configured model/default", "Use", "Activation"],
        [
            ("OpenAI", "gpt-4o-mini for chat and intent", "Narration, intent refinement, titles, optional agent planning", "OPENAI_API_KEY"),
            ("Google Gemini", "gemini-1.5-flash; report setting gemini-1.5-pro", "Fallback narration and semantic/query refinement", "GEMINI_API_KEY or GOOGLE_API_KEY"),
            ("Anthropic", "claude-sonnet-4-6 in config.py", "Fallback narration", "ANTHROPIC_API_KEY"),
            ("DeepSeek", "deepseek-chat", "Intent parsing through OpenAI-compatible API", "DEEPSEEK_API_KEY"),
            ("DeepAnalyze", "deepanalyze-8b; local fallback label phi3:mini", "Remote/local OpenAI-compatible reasoning and narration", "DeepAnalyze API or local base URL"),
            ("Deterministic", "Rule/template code", "All calculations plus fallback summaries and answers", "Always available"),
        ],
        font_size=7.8,
    )
    b.heading("11.1 Provider order", 2)
    b.paragraph("In auto mode, llm_provider.py tries configured OpenAI, Gemini, Anthropic, and DeepAnalyze providers in that order. A specifically requested provider is placed first. Errors are recorded by provider type, and a missing/failed chain returns None so deterministic narration is used.")
    b.heading("11.2 Agentic chat loop", 2)
    b.paragraph("When an LLM is configured, AgentLoop can choose deterministic tools, observe their computed JSON, and produce an answer restricted to those observations. The UI can display the thought/tool/argument/observation trace. Without an LLM, SessionService uses the deterministic answer path.")
    b.callout("Configuration mismatch", "config.py defaults CLAUDE_MODEL to claude-sonnet-4-6, while dataverse_backend/.env.example currently lists claude-3-5-sonnet-20241022. Deployment behavior follows the environment value when supplied, so documentation/configuration should be aligned.", kind="warning")
    b.evidence("dataverse_backend/app/services/llm_provider.py", "llm_service.py", "agent_loop.py", "query_planner.py", "semantic_mapper.py", "report_narrator.py", "title_generator.py", "app/core/config.py")

    b.page_break()
    b.heading("12. Supabase, authentication, database, and storage", 1)
    b.heading("12.1 Authentication", 2)
    b.table(
        "Authentication flow and security rules",
        ["Operation", "Implementation"],
        [
            ("Signup", "FastAPI validates input, then POSTs email/password/user metadata to Supabase Auth /auth/v1/signup"),
            ("Password policy", "At least 12 characters; uppercase, lowercase, digit, special character; cannot contain email local-part"),
            ("Password storage", "Supabase Auth owns hashing/storage; the application database does not store password fields"),
            ("Email confirmation", "Signup is rejected operationally if Supabase returns an immediate access token; confirmation link redirects to /login?confirmed=true"),
            ("Resend", "Supabase /auth/v1/resend with type=signup"),
            ("Login", "Supabase password grant returns access token, refresh token, expiry, and user"),
            ("Refresh", "Supabase refresh_token grant"),
            ("Token verification", "Legacy HS256 checked locally when JWT secret exists; asymmetric/new tokens checked against Supabase /auth/v1/user"),
            ("Verification cache", "Validated token-to-user mapping cached in memory for five minutes"),
        ],
        font_size=7.8,
    )

    b.heading("12.2 PostgreSQL metadata schema", 2)
    b.table(
        "Supabase application tables",
        ["Table", "Important fields", "Purpose"],
        [
            ("chat_sessions", "id, user_id, title, status, active_dataset_id, metadata, timestamps", "Ownership and conversation/workspace state"),
            ("chat_messages", "id, session_id, role, content, message_type, payload, created_at", "User, assistant, system, and agent messages"),
            ("datasets", "id, session_id, user_id, filename, storage_path, file size, shape, columns, profile, semantic map, status", "Dataset metadata and object location"),
            ("agent_runs", "id, session_id, dataset_id, agent_name, status, input, output, error, timestamps", "Execution trace persistence"),
            ("reports", "id, session_id, dataset_id, title, type, format, storage_path, public_url, metadata", "Generated report metadata"),
            ("auth.users", "Managed by Supabase", "Credentials, confirmation state, and user identity"),
        ],
        font_size=7.6,
    )

    b.heading("12.3 Object storage", 2)
    b.bullets([
        "dataverse-datasets is a private bucket for uploaded dataset objects.",
        "dataverse-reports is a private bucket for generated HTML and PDF reports.",
        "The backend generates signed download URLs with a default one-hour lifetime.",
        "SUPABASE_SERVICE_ROLE_KEY remains backend-only and is used for REST and Storage operations.",
        "If Supabase is not configured, LocalPersistence stores table snapshots and files under session_storage/dataverse_chat.",
    ])
    b.callout("Schema caveat", "supabase/migrations/001_dataverse_schema.sql enables RLS but defines no browser policies because the backend service role bypasses RLS. dataverse_backend/supabase_schema.sql contains a similar schema with optional policies commented out. These duplicate schema sources should be consolidated before production migration management.", kind="warning")
    b.evidence("supabase/migrations/001_dataverse_schema.sql", "dataverse_backend/supabase_schema.sql", "dataverse_backend/app/services/auth_service.py", "supabase_client.py", "session_service.py")

    b.heading("13. Security assessment", 1)
    b.table(
        "Implemented security controls",
        ["Control", "Implementation", "Assessment"],
        [
            ("Server-side password policy", "12-character complexity and email-name exclusion", "Strong baseline"),
            ("Email verification", "Required Supabase confirmation link", "Prevents immediate unverified access when dashboard configuration is correct"),
            ("No plaintext password storage", "Credentials sent to Supabase Auth only", "Correct responsibility boundary"),
            ("JWT validation", "HS256 local verification or Supabase user endpoint", "Supports key-algorithm transition"),
            ("IDOR guard", "Session owner must match resolved identity", "Protects owned session routes"),
            ("Rate limiting", "In-memory sliding window on sensitive POST routes", "Useful for one process; not distributed"),
            ("Security headers", "Frame denial, MIME sniff prevention, referrer and permissions policy", "Good browser baseline"),
            ("Private buckets", "Datasets and reports are non-public", "Correct default"),
            ("Service-role isolation", "Service key is backend-only", "Required for current persistence design"),
            ("Upload checks", "Size, extension, non-empty dataframe, domain guard", "Reduces obvious misuse"),
            ("Production error masking", "Internal exception details hidden outside development", "Reduces information leakage"),
        ],
        font_size=7.5,
    )
    b.table(
        "Security and production-hardening gaps",
        ["Priority", "Gap", "Risk", "Recommended change"],
        [
            ("High", "Tokens stored in localStorage", "An XSS vulnerability could read access and refresh tokens", "Prefer secure, HttpOnly, SameSite cookies or a hardened token strategy plus strict CSP"),
            ("High", "X-Dataverse-User accepted without bearer token", "Caller-selected workspace identity can be spoofed", "Require verified Supabase identity for all non-public data routes; remove anonymous identity fallback"),
            ("High", "Legacy sessions without owner remain open", "Known session IDs may be accessible", "Migrate/assign owners and deny ownerless sessions in production"),
            ("High", "Progress SSE route does not call _authorize", "Session progress may be observable by ID", "Verify bearer identity and session ownership before streaming"),
            ("High", "Service role bypasses RLS", "Backend authorization errors expose all rows", "Keep strict application checks and add user-scoped defense-in-depth RPC/policies where feasible"),
            ("Medium", "In-memory rate limiter/cache", "Resets on restart and is not shared across workers", "Use Redis or another shared store in multi-instance deployment"),
            ("Medium", "Extension validation only", "File content may not match extension", "Add MIME/signature checks, workbook limits, decompression limits, and malware scanning"),
            ("Medium", "Two Supabase schema files differ", "Migration drift", "Choose one migration source and version every change"),
            ("Medium", "CORS includes HTTP localhost defaults", "Suitable for development, not production", "Use HTTPS production origins only"),
            ("Medium", "No explicit Content-Security-Policy", "Reduced protection against script injection", "Add a tested CSP in Next.js/FastAPI headers"),
        ],
        font_size=7.2,
    )

    b.heading("14. Reporting, charts, and user-visible outputs", 1)
    b.table(
        "Output types",
        ["Output", "Produced by", "Content"],
        [
            ("Chat answer", "SessionService / AnalysisPipeline", "Answer, summary, warnings, recommendations, next questions"),
            ("KPI cards", "business_metrics.py", "Revenue, quantity, profit, margin, order value, counts, provenance"),
            ("Tables", "business/product/food/financial analysis", "Rankings, frequencies, metrics, audit-friendly rows"),
            ("Charts", "data_quality.py and business modules", "Bar, line, pie/donut, histogram, scatter, boxplot, feature importance, confusion matrix"),
            ("Agent trace", "agent_loop.py", "Thought, tool, arguments, deterministic observation"),
            ("Quality diagnosis", "quality_doctor.py", "Issue list, fixes, before/after"),
            ("Certificate", "certificate.py", "Algorithm, data/result fingerprints, verified number count"),
            ("HTML report", "report_generator.py", "Self-contained styled narrative, SVG/Chart.js-ready visuals, tables, evidence"),
            ("PDF report", "report_generator.py", "ReportLab-rendered static report with fallback manual PDF writer"),
        ],
        font_size=7.6,
    )
    b.paragraph("Chart specifications are normalized and validated before delivery. Invalid charts are dropped with warnings. ReportComposer tracks semantic fingerprints to prevent repeated charts, tables, or insights and adds data-specific explanations and takeaways.")
    b.evidence("dataverse_backend/app/services/data_quality.py", "report_composer.py", "report_generator.py", "frontend/components/dashboard/*")

    b.page_break()
    b.heading("15. Deployment and configuration", 1)
    b.heading("15.1 Local development", 2)
    b.bullets([
        "npm run dev launches both the FastAPI backend and the Next.js frontend.",
        "FastAPI listens on http://127.0.0.1:8000 with reload enabled.",
        "Next.js listens on https://127.0.0.1:3000 using local certificate files.",
        "Next.js proxies /backend/* to the HTTP backend so browser requests remain same-origin HTTPS.",
    ])
    b.heading("15.2 Containers", 2)
    b.paragraph("docker-compose.yml defines backend port 8000 and frontend port 3000. The backend image uses python:3.12-slim, installs requirements-mvp.txt, copies the application, and exposes an HTTP health check. The frontend image builds and starts Next.js. Current compose examples use HTTP frontend/backend URLs and should be fronted by a TLS reverse proxy in production.")
    b.heading("15.3 Environment configuration", 2)
    b.table("Environment-variable categories", ["Category", "Variables", "Purpose"], ENV_ROWS, font_size=6.9)
    b.callout("Dependency reproducibility", "The backend Dockerfile installs requirements-mvp.txt, whose packages are largely unpinned. requirements-full.txt contains many pins but is not the Docker input. Freeze the deployed requirements and remove duplicated/unneeded entries for reproducible production images.", kind="warning")
    b.evidence("frontend/scripts/dev.mjs", "frontend/next.config.ts", "docker-compose.yml", "dataverse_backend/Dockerfile", "frontend/Dockerfile", "*.env.example")

    b.heading("16. Complete library and dependency inventory", 1)
    b.paragraph("The table records versions installed in the audited environment where available. Declared manifests are not fully consistent: several backend packages are unpinned in the active MVP requirements, and installed reportlab, joblib, and pytest-asyncio versions differ from requirements-full.txt.")
    b.table("Important frontend and backend dependencies", ["Library", "Audited version", "Use status", "Role"], DEPENDENCY_ROWS, font_size=7.0)
    b.heading("16.1 Declared but not part of the active core flow", 2)
    b.table(
        "Supporting or inactive dependencies/configuration",
        ["Item", "Repository status", "Interpretation"],
        [
            ("SQLAlchemy / aiosqlite / DATABASE_URL", "Declared in full requirements/settings; active session flow uses Supabase REST/local files", "Not required for current UI flow"),
            ("Redis / Celery", "Settings exist; no active worker pipeline imported", "Future/legacy configuration"),
            ("MinIO / AWS S3", "Settings exist; SessionService uses Supabase/local persistence", "Inactive alternatives"),
            ("Plotly / Kaleido / Matplotlib", "Declared full stack; application reports use custom SVG/ReportLab; matplotlib appears through environment/tests", "Not central to runtime chart delivery"),
            ("Supabase Python package", "Declared but not installed; no direct import", "Custom HTTPX client is the actual implementation"),
            ("requests", "Declared; active external calls use HTTPX", "Currently unnecessary for active app code"),
            ("Sentry settings", "Configuration fields exist; no Sentry initialization found", "Not operational"),
        ],
        font_size=7.8,
    )

    b.heading("17. Test and build verification", 1)
    b.table(
        "Verification performed for this report",
        ["Command", "Result", "Evidence"],
        [
            (".venv/Scripts/python -m pytest -q", "PASS: 214 tests; 14 dependency deprecation warnings; 140.50 seconds", "Backend API, auth security, analytics, agents, reports, XAI, quality, root cause, what-if, uploads"),
            ("npm run lint", "PASS", "ESLint completed without errors"),
            ("npm run build", "PASS", "Next.js compiled, type-checked, and generated 10 static pages/routes"),
            ("Dataset audit", "PASS", "35,000 x 21; zero missing cells; zero duplicate rows"),
            ("Model artifact inspection", "PASS", "Three sklearn Pipeline objects with preprocess/model steps and expected estimators"),
        ],
        font_size=8,
    )
    b.paragraph("The production build reports 102 kB shared first-load JavaScript. The dashboard route is approximately 102 kB route size and 208 kB first-load JavaScript, while public pages are about 109-110 kB first load.")
    b.callout("Warnings", "The 14 pytest warnings come from Matplotlib/PyParsing deprecations in installed dependencies, not failing application assertions. The Next build output says 'Skipping linting', but lint was run independently and passed.", kind="info")

    b.page_break()
    b.heading("18. Limitations, uncertainty, and robustness", 1)
    b.table(
        "Technical limitations that affect interpretation",
        ["Area", "Limitation", "Effect"],
        [
            ("Dataset provenance", "No original source description, license, collection period, sampling method, or encoding map", "Business/generalization claims cannot be established"),
            ("Saved total-sales model", "Engineered features reconstruct the target", "R2 = 1.0 is leakage, not forecasting evidence"),
            ("Saved stockout model", "Positive recall 26.3% and precision 19.9%", "Misses most actual stockouts and generates many false alarms"),
            ("Model validation", "Single 75/25 random holdout; no cross-validation, temporal split, external set, calibration, or uncertainty intervals", "Metrics may be unstable or optimistic"),
            ("Live feature leakage", "Safety check is mainly name/identifier based", "Derived or hidden target proxies may still enter runtime models"),
            ("Causal claims", "Root cause and what-if are deterministic descriptive calculations", "They do not establish causal effects"),
            ("LLM behavior", "Providers are optional and external output can vary", "Narrative must remain bounded to computed facts"),
            ("Persistence", "Local fallback can hide missing Supabase configuration", "A locally working app may not be production-ready"),
            ("Security", "Token storage and anonymous identity fallbacks remain", "Production hardening is still required"),
            ("Documentation", "Historical docs describe routes/features that have changed", "Code and current tests must remain the source of truth"),
        ],
        font_size=7.3,
    )

    b.heading("19. Recommended next steps", 1)
    b.bullets([
        "Remove target-derived features from each saved training task, retrain, and use cross-validation plus a time-aware or external holdout where the business problem is temporal.",
        "For stockout detection, optimize positive-class recall/precision, tune thresholds, report PR-AUC/ROC-AUC, calibrate probabilities, and validate on new stores/time periods.",
        "Create a versioned dataset card containing origin, license, collection dates, feature definitions, categorical encoding maps, intended use, and prohibited use.",
        "Require Supabase-authenticated identities for all dataset/session/report routes and remove X-Dataverse-User as an authorization mechanism.",
        "Protect refresh/access tokens with HttpOnly secure cookies or an equivalent hardened browser session design, then add CSP and CSRF controls appropriate to that design.",
        "Authorize the progress SSE route and deny access to ownerless legacy sessions.",
        "Consolidate the two Supabase schemas into one ordered migration history and add defense-in-depth RLS policies where direct browser access is expected.",
        "Pin requirements used by Docker, remove unused dependencies, and align .env.example defaults with config.py.",
        "Add integration tests against a disposable Supabase project for confirmation email, JWT algorithms, RLS, private bucket access, and signed URLs.",
        "Either integrate the saved retail artifacts as an explicit versioned inference feature or label/archive them as offline evaluation outputs to avoid architectural confusion.",
    ], ordered=True)

    b.heading("20. Final implementation conclusion", 1)
    b.paragraph("The repository implements a substantial deterministic analytics application: real Supabase email/password authentication, a typed Next.js interface, a FastAPI session API, semantic dataset analysis, business KPIs, data-quality checks, runtime AutoML candidates, XAI, counterfactuals, root-cause decomposition, reproducibility receipts, and HTML/PDF reporting. The test and build results support functional completeness at the repository level.")
    b.paragraph("The strongest engineering characteristic is separation of numerical computation from optional language-model narration. The main areas preventing a production-readiness claim are saved-model validation quality, incomplete dataset provenance, token/anonymous identity security, duplicate persistence schemas, and unpinned deployment dependencies.")

    b.page_break()
    b.heading("Appendix A. Backend service-module inventory", 1)
    b.table("Backend service modules and status", ["Module", "Status", "Responsibility"], SERVICE_ROWS, font_size=6.9)
    b.evidence("dataverse_backend/app/services/*.py")

    b.page_break()
    b.heading("Appendix B. Model metadata detail", 1)
    for target in ("total_sales", "profit", "stockout_flag"):
        entry = metadata[target]
        b.heading(f"B.{1 + ['total_sales', 'profit', 'stockout_flag'].index(target)} {target}", 2)
        metric_rows = [(key, value) for key, value in entry["metrics"].items()]
        b.table(f"{target} saved metrics", ["Metric", "Value"], metric_rows, font_size=8.3)
        b.table(
            f"{target} top recorded feature importance",
            ["Rank", "Feature", "Normalized importance"],
            [(idx + 1, item["feature"], f"{item['importance']:.6f}") for idx, item in enumerate(entry["feature_importance"])],
            font_size=8,
        )
        if target == "stockout_flag":
            b.table(
                "stockout_flag confusion-matrix cells",
                ["Actual", "Predicted", "Count"],
                [(item["actual"], item["predicted"], f"{item['count']:,}") for item in entry["confusion_matrix"]],
                font_size=8.3,
            )

    b.page_break()
    b.heading("Appendix C. Evidence and source map", 1)
    b.table(
        "Primary report evidence",
        ["Subject", "Repository evidence"],
        [
            ("Architecture", "docs/markdown/02_architecture_and_design/ARCHITECTURE.md; app/main.py; session_service.py; frontend/lib/dataverse-api.ts"),
            ("Authentication", "auth_service.py; auth_routes.py; frontend/lib/auth.tsx; tests/test_auth_security.py"),
            ("Dataset ingestion", "DatasetAgent; upload_parsing.py; data_loader.py; session_store.py"),
            ("Analytics", "analysis_pipeline.py; business_metrics.py; data_quality.py; data_profiler.py"),
            ("Runtime models", "modeling.py; target_inference.py; xai.py; counterfactual.py"),
            ("Saved models", "scripts/train_retail_mart.py; models/model_metadata.json; models/*.pkl"),
            ("Retail data", "data/retail_mart_processed_v1.csv"),
            ("Supabase", "supabase/migrations/001_dataverse_schema.sql; supabase_client.py; session_service.py"),
            ("Reports", "report_composer.py; report_generator.py; report_narrator.py"),
            ("Frontend", "frontend/app/*; frontend/components/*; frontend/lib/*; frontend/package*.json"),
            ("Deployment", "Dockerfiles; docker-compose.yml; scripts/dev.mjs; .env.example files"),
            ("Verification", "dataverse_backend/tests/*; npm lint/build output; model/dataset audit commands"),
        ],
        font_size=7.6,
    )

    b.page_break()
    b.heading("Appendix D. Complete current-folder file inventory", 1)
    b.paragraph(
        "This appendix itemizes every safely reportable file found in the current FINAL3 folder at generation time. The purpose column explains the file's repository role. Environment files are listed by name only; their contents are not included. Installed dependency internals, private session data, caches, Git objects, temporary files, and certificate/key material are accounted for separately in Appendix E."
    )
    grouped_inventory: dict[str, list[dict[str, object]]] = {}
    for item in inventory_files:
        grouped_inventory.setdefault(str(item["category"]), []).append(item)
    for category in sorted(grouped_inventory, key=lambda value: (CATEGORY_ORDER.get(value, 99), value)):
        items = grouped_inventory[category]
        b.heading(category, 2)
        b.paragraph(f"{len(items):,} files in this responsibility group.")
        b.table(
            f"File inventory - {category}",
            ["File or relative path", "Type", "Size / lines", "Source state", "Purpose"],
            [
                (
                    item["path"],
                    item["type"],
                    f"{format_bytes(int(item['bytes']))} / {item['lines'] if item['lines'] != '' else 'binary'}",
                    item["source_control"],
                    item["role"],
                )
                for item in items
            ],
            font_size=5.5,
        )

    b.page_break()
    b.heading("Appendix E. Summarized dependency, cache, private, and runtime directories", 1)
    b.paragraph(
        "These directories are part of the physical folder and are included in the total file accounting. Their internal files are not reproduced because they are third-party package installations, generated caches, Git internals, mutable private user/session data, temporary verification material, or local certificate/key material. Their authoritative project-facing definitions are the manifests, lockfiles, build results, and configuration files documented elsewhere in this report. The companion DataVerse_AI_Current_Folder_File_Inventory.csv and DataVerse_AI_Current_Folder_Audit.json files preserve the searchable rows and aggregate accounting."
    )
    b.table(
        "Safely summarized directories",
        ["Directory", "Files", "Size", "Reason for aggregate treatment"],
        [
            (
                item["path"],
                f"{int(item['files']):,}",
                format_bytes(int(item["bytes"])),
                item["reason"],
            )
            for item in excluded_directories
        ],
        font_size=6.8,
    )
    extension_rows = sorted(
        inventory_summary["extensions"].items(),
        key=lambda item: (-item[1], item[0]),
    )
    b.table(
        "Safely itemized files by extension/type",
        ["Extension or type", "Files", "Meaning"],
        [
            (
                extension,
                f"{count:,}",
                "Exact paths and purposes are listed in Appendix D.",
            )
            for extension, count in extension_rows
        ],
        font_size=7.2,
    )
    b.save(DOCX_OUT)

    source_notes = ROOT / "docs" / "markdown" / "08_generated_notes" / "DATAVERSE_AI_TECHNICAL_REPORT_SOURCE_NOTES.md"
    source_notes.parent.mkdir(parents=True, exist_ok=True)
    source_notes.write_text(
        """# DataVerse AI technical report source notes

Audit date: 2026-07-14

## Report spine

- Question: What is present in the complete current FINAL3 folder, what is each project file for, and which models, datasets, libraries, components, services, security controls, APIs, artifacts, and deployment mechanisms are actually used?
- Answer: Deterministic Next.js/FastAPI/Pandas/scikit-learn system with conditional XAI/LLMs and Supabase-backed auth/persistence.
- Key caveats: saved model artifacts are not loaded by live app; total_sales leakage; weak minority stockout performance; external Supabase production state not verified.

## Folder inventory scope

- Every safely reportable project and generated file is included in the CSV/JSON inventory and Appendix D.
- Third-party installations, Git internals, build/test caches, browser traces, session storage, temporary conversions, and local certificate material are counted in Appendix E but their internal contents are not reproduced.
- Environment files are inventoried by filename only; no secret values are read into or written to the report.

## Chart map

1. System architecture: component/flow diagram; source app code and existing architecture asset; supports the end-to-end runtime claim.
2. Saved model feature importance: three horizontal bar panels; source models/model_metadata.json; supports leakage and driver interpretation; single-root palette per panel; exact values shown.
3. Stockout confusion matrix: 2x2 matrix; source models/model_metadata.json; supports minority-class performance limitation; exact counts and test-share labels shown.

## Verification

- Backend: 214 passed, 14 dependency deprecation warnings.
- Frontend lint: passed.
- Frontend production build: passed.
- PDF visual verification is performed after Word export.

## Omitted visuals

- No cross-model performance chart because regression R2/RMSE/MAE and classification accuracy/F1 are not directly comparable.
- Dependency, API, and database inventories use tables because exact lookup is the reader's primary task.
""",
        encoding="utf-8",
    )
    print(f"docx={DOCX_OUT}")
    print(f"source_notes={source_notes}")
    print(f"inventory_csv={inventory['csv_path']}")
    print(f"inventory_json={inventory['json_path']}")
    print(f"figures={feature_path};{confusion_path}")


if __name__ == "__main__":
    build_report()
