"""Generate the DataVerse AI report corrections + new-chapter Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

VIOLET = RGBColor(0x5B, 0x21, 0xB6)
INK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x47, 0x55, 0x69)
GREEN = RGBColor(0x05, 0x96, 0x69)

doc = Document()
st = doc.styles["Normal"].font
st.name = "Calibri"; st.size = Pt(11)


def h(txt, level=1, color=VIOLET):
    p = doc.add_heading(txt, level=level)
    for r in p.runs:
        r.font.color.rgb = color
        r.font.name = "Cambria"
    return p


def para(txt, size=11, italic=False, bold=False, color=INK, space=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(txt); r.font.size = Pt(size); r.italic = italic; r.bold = bold; r.font.color.rgb = color
    return p


def bullet(txt, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if lead:
        r = p.add_run(lead + " "); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = INK
    r2 = p.add_run(txt); r2.font.size = Pt(11); r2.font.color.rgb = MUTED
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hd in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        run = c.paragraphs[0].add_run(hd); run.bold = True; run.font.size = Pt(10.5)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(val); run.font.size = Pt(10)
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    return t


# ---------------- TITLE ----------------
title = doc.add_heading("DataVerse AI — Report Corrections & New Chapter", level=0)
for r in title.runs:
    r.font.color.rgb = INK; r.font.name = "Cambria"
para("How to use this document: Section 1 lists global find-and-replace rules to apply across the report. "
     "Section 2 gives exact rewrites for the densest passages. Section 3 is a ready-to-paste new chapter documenting "
     "the verifiability features. Section 4 updates the results with measured verification data. "
     "Apply Section 1 first, then Section 2.", italic=True, color=MUTED)
para("Why these changes: the current report describes several technologies the built system does not use "
     "(LangChain, PyCaret, Plotly/Matplotlib, LIME, Llama) and frames the system as a generic “agent-based / "
     "multi-agent” platform. The actual implementation is a deterministic, strictly two-agent pipeline whose "
     "differentiator — verifiable, reproducible numbers — is currently undocumented.", color=MUTED)

# ---------------- SECTION 1: GLOBAL RULES ----------------
h("1. Global Find-and-Replace Rules", 1)
para("Apply these throughout the report. Important: in the Literature Review (Chapter 2), references to SHAP, LIME, "
     "and other tools describe prior/related work — leave those unchanged. Only change text that describes "
     "DataVerse AI’s own design, features, or stack.", italic=True, color=MUTED)
table(
    ["Find (in DataVerse’s own descriptions)", "Replace with", "Reason"],
    [
        ["“SHAP and LIME”, “SHAP or LIME”, “SHAP/LIME”", "“SHAP”", "The system uses SHAP only (with a feature-importance fallback). LIME is not implemented."],
        ["“PyCaret” (and “Scikit-learn and PyCaret”)", "“scikit-learn”", "Modelling uses scikit-learn only (Ridge, RandomForest). PyCaret is not used."],
        ["“Plotly and Matplotlib”, “Matplotlib”, “Plotly” (for the app’s charts)", "“hand-rolled SVG charts (ReportLab for PDF)”", "The live pipeline renders charts as lightweight SVG; ReportLab builds the PDF. Plotly/Matplotlib are not in the live path."],
        ["“OpenAI and LangChain”, “OpenAI/LangChain”, “OpenAI API, Llama, and LangChain”", "“OpenAI / Google Generative AI (optional)”", "The LLM layer is optional and used only to polish narration; there is no LangChain or Llama. With no keys, it runs in offline Mock mode."],
        ["“agent-based platform”, “multi-agent system”", "“two-agent system (DatasetAgent + AnalystAgent)”", "The architecture is exactly two agents by design."],
        ["“the LLM agent executes / decides the analysis”", "“the LLM only polishes narration; all numbers are computed deterministically”", "Deterministic-first: the LLM never produces a number."],
    ],
    widths=[2.6, 2.3, 2.4],
)

# ---------------- SECTION 2: KEY PASSAGES ----------------
h("2. Key Passage Rewrites (exact replacements)", 1)

h("2.1  Chapter 1 §1.1 — Introduction (page 5)", 2)
para("CURRENT:", bold=True, color=MUTED, space=2)
para("“The user interface is built using Next.js, and the back end is powered by FastAPI… OpenAI and Lang Chain let "
     "users talk to the system naturally… The main part of the system uses Pandas, PyCaret, and Scikit-learn… Charts "
     "are made with Plotly and Matplotlib, and tools like SHAP and LIME explain how the models work…”", italic=True, color=MUTED)
para("REPLACE WITH:", bold=True, color=GREEN, space=2)
para("The user interface is built using Next.js (React), and the back end is powered by FastAPI. The system is "
     "deterministic-first: all metrics, exploratory analysis, and predictions are computed in Pandas and "
     "scikit-learn, while an optional LLM (OpenAI or Google Generative AI) is used only to polish the wording of the "
     "already-computed results — it never produces a number. Predictive models (Ridge regression and Random Forest) "
     "are explained using SHAP. Charts are rendered as lightweight SVG and reports as HTML/PDF (ReportLab). Crucially, "
     "every number the system reports is accompanied by a verifiable “receipt” and a reproducibility certificate, so "
     "results can be independently re-checked and trusted. With no API keys configured, the entire pipeline still "
     "runs in an offline “Mock mode.”")

h("2.2  Chapter 2 — Literature Review (pages 9–10): DataVerse-specific sentences", 2)
para("Two sentences in Chapter 2 describe DataVerse itself (not prior work) and must be corrected:", color=MUTED)
para("• “It employs an LLM (through OpenAI/LangChain) to understand user intent and assigns the actual analytical "
     "tasks to fixed Python modules… such as Pandas, NumPy, and Plotly” →", bold=True, color=MUTED, space=2)
para("“DataVerse AI uses an optional LLM (OpenAI or Google Generative AI) only to interpret phrasing and polish "
     "narration; all analytical work is performed by deterministic Python modules (Pandas, NumPy, scikit-learn), and "
     "visualisations are rendered as SVG.”", italic=True)
para("• “…underpin DataVerse AI’s explainability approach, combining SHAP and LIME for detailed model interpretation” →",
     bold=True, color=MUTED, space=2)
para("“…underpin DataVerse AI’s explainability approach, using SHAP for model interpretation (with a "
     "feature-importance fallback).”", italic=True)

h("2.3  Chapter 4 §4.2 — Methodology (page 27)", 2)
para("CURRENT:", bold=True, color=MUTED, space=2)
para("“The analytics engine, built with Python, uses Pandas, Scikit-learn, PyCaret, SHAP, and Matplotlib. The LLM "
     "engine relies on the OpenAI API, Llama, and Lang Chain.”", italic=True, color=MUTED)
para("REPLACE WITH:", bold=True, color=GREEN, space=2)
para("The analytics engine, built with Python, uses Pandas, NumPy, and scikit-learn (Ridge, Random Forest), with SHAP "
     "for explainability and lightweight SVG charts (ReportLab for PDF). The optional LLM layer integrates OpenAI or "
     "Google Generative AI for narration polish only and is fully optional (offline Mock mode otherwise).")

h("2.4  Chapter 5 §5.2.1 — Architectural Design (page 30)", 2)
para("CURRENT:", bold=True, color=MUTED, space=2)
para("“…Pandas and NumPy handle data processing, Plotly and Matplotlib generate visual outputs, Scikit-learn and "
     "PyCaret manage predictive modeling, and SHAP and LIME provide model explanations.”", italic=True, color=MUTED)
para("REPLACE WITH:", bold=True, color=GREEN, space=2)
para("Pandas and NumPy handle data processing, hand-rolled SVG charts (and ReportLab for PDF) generate visual "
     "outputs, scikit-learn manages predictive modelling (Ridge, Random Forest), and SHAP provides model "
     "explanations. The architecture is organised around exactly two agents — the DatasetAgent and the AnalystAgent — "
     "coordinating a deterministic analysis pipeline; the optional LLM only narrates the computed results.")
para("Also change, earlier on page 30: “…layered and modular architecture supported by agent-based orchestration” → "
     "“…layered and modular architecture built on exactly two agents (DatasetAgent and AnalystAgent) coordinating a "
     "deterministic pipeline.”", color=MUTED)

h("2.5  Chapter 6 §6.8 — Tools table (page 44)", 2)
para("Replace the “Matplotlib — Static chart generation” row with: "
     "“SVG / ReportLab — Chart rendering and PDF generation”. Remove any PyCaret/LangChain/Llama rows if present, and "
     "ensure the LLM row reads “OpenAI / Google Generative AI (optional, narration only)”.", color=MUTED)

# ---------------- SECTION 3: NEW CHAPTER ----------------
doc.add_page_break()
h("3. New Chapter (ready to paste): Verifiable AI — Trust & Reproducibility", 1)
para("Suggested placement: as a new section at the end of Chapter 5 (Design) or as a short standalone chapter before "
     "Results. It documents the project’s core differentiator.", italic=True, color=MUTED)

h("Verifiable AI: Trust and Reproducibility", 2, INK)
para("A defining limitation of LLM-based analytics tools is that their numerical outputs can be fabricated by the "
     "language model and cannot be independently reproduced. DataVerse AI addresses this directly: it is built on a "
     "deterministic-first principle, and exposes a layered set of trust features that make every reported number "
     "transparent, auditable, and re-verifiable. Two of these features are, to the best of our knowledge, not offered "
     "by existing conversational analytics tools.")

h("3.1  Deterministic-first computation", 3, INK)
para("Every metric, statistic, trend, and prediction shown to the user is computed in Pandas or scikit-learn. The "
     "optional Large Language Model is restricted to polishing the wording of these pre-computed results; it can "
     "never originate or alter a value. This guarantee is the foundation on which the remaining trust features rest.")

h("3.2  “Show-the-Math” provenance receipts", 3, INK)
para("Each business KPI is accompanied by a provenance receipt that records the operation performed (e.g. SUM, MEAN), "
     "a plain-English formula, the source column(s), the number of contributing rows, and a sample of the actual rows "
     "that produced the value. A regression test asserts that each receipt’s value equals the reported metric, so the "
     "guarantee is enforced automatically in continuous integration.")

h("3.3  Full audit trail", 3, INK)
para("The same provenance mechanism is extended to exploratory statistics, correlations, trends, and model-evaluation "
     "metrics, producing a single audit trail of receipts that can be downloaded as a JSON log for record-keeping.")

h("3.4  Reproducibility certificate (novel)", 3, INK)
para("Each analysis is issued a tamper-evident certificate consisting of two SHA-256 fingerprints: one of the raw "
     "dataset and one of the deterministic results. A one-click “re-verify” operation re-runs the deterministic "
     "pipeline on the raw data and confirms that every number reproduces exactly, while also detecting whether the "
     "data or the results were altered. Because LLM-orchestrated tools do not produce reproducible output, they have "
     "nothing comparable to certify — making this capability a genuine differentiator.")

h("3.5  Verified what-if simulator (novel)", 3, INK)
para("Users can adjust a numeric lever (for example, “+10% price”) and have the business KPIs recomputed on the "
     "modified data through the same deterministic pipeline. Each hypothetical KPI carries its own provenance "
     "receipt, so a what-if scenario is as verifiable as the original analysis — in contrast to LLM-generated "
     "estimates, which are neither reproducible nor explainable.")

h("3.6  Data Quality Doctor", 3, INK)
para("The system automatically detects data-quality issues (duplicate rows, missing values, constant columns, and "
     "numeric values stored as text) and proposes one-click, deterministic fixes (median/mode imputation, column "
     "removal, type casting, de-duplication), each with a before/after impact. Applying the fixes produces a cleaned "
     "dataset and a fresh, verifiable analysis.")

h("3.7  Verified report export & conversational drill-downs", 3, INK)
para("Analyses can be exported as a self-contained HTML report in which every KPI embeds its receipt and the "
     "reproducibility certificate, allowing a recipient who holds the data to re-verify the figures independently. "
     "Within the chat interface, clicking any number issues a verified follow-up query, turning the workspace into an "
     "investigative, fully traceable conversation.")

# ---------------- SECTION 4: RESULTS ADDENDUM ----------------
doc.add_page_break()
h("4. Results Addendum (measured verification data)", 1)
para("Add the following to Chapter 7 (Results & Discussion), e.g. as “7.x Engineering Verification.”", italic=True, color=MUTED)
table(
    ["Verification check", "Result"],
    [
        ["Backend automated tests (pytest)", "129 tests passing, including the trust test asserting every receipt equals its metric"],
        ["Live reproducibility re-verification", "20 / 20 deterministic numbers reproduced exactly from the raw data"],
        ["Frontend quality gates", "ESLint clean; production build (Next.js) succeeds; verified in a real browser with no console errors"],
        ["Database insert", "< 100 ms"],
        ["Full 35,000-row analysis", "< 500 ms"],
        ["Offline operation", "Full pipeline runs in Mock mode with zero API keys"],
    ],
    widths=[3.2, 4.1],
)
h("4.1  Comparison with existing tools", 2)
table(
    ["Capability", "ChatGPT ADA", "Julius AI", "Tableau", "DataVerse AI"],
    [
        ["Natural-language Q&A", "Yes", "Yes", "No", "Yes"],
        ["Deterministic numbers", "No", "No", "Yes", "Yes"],
        ["Per-number receipts", "No", "No", "No", "Yes"],
        ["Reproducibility certificate", "No", "No", "No", "Yes"],
        ["Verified what-if", "No", "No", "No", "Yes"],
        ["One-click data cleaning", "Partial", "Partial", "No", "Yes"],
    ],
    widths=[2.5, 1.2, 1.1, 1.0, 1.5],
)
para("")
para("Note on dataset/results: the measured KPIs in Chapter 7 (Total Sales $2,248,662.62; Total Profit $336,938.93; "
     "Profit Margin 14.98%; 69,546 units; Region 2 best; Category 1 best; Store 27 top) remain correct and need no "
     "change — they are produced by the deterministic pipeline described above.", color=MUTED)

OUT = r"C:/Users/mouav/Downloads/DataVerse_Report_Corrections.docx"
doc.save(OUT)
print("saved", OUT)
