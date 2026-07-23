from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "pdf" / "DataVerse_AI_Final_Report_Simple_Architecture_Formatted_2026.docx"
OUTPUT = ROOT / "output" / "pdf" / "DataVerse_AI_Final_Report_Latest_Architecture_2026.docx"
DIAGRAM = ROOT / "output" / "architecture" / "DataVerse_AI_System_Architecture_Updated.png"
FONT_NAME = "Times New Roman"


def set_run_font(run, size: float = 12, *, bold: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), FONT_NAME)


def replace_text(paragraph, text: str, *, centered: bool = False, bold: bool = False) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    set_run_font(run, 12, bold=bold)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY


def replace_picture(paragraph, image_path: Path) -> None:
    for child in list(paragraph._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            paragraph._p.remove(child)
    paragraph.add_run().add_picture(str(image_path), width=Inches(6.6))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def find_paragraph(document, exact_text: str):
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == exact_text:
            return index, paragraph
    raise KeyError(f"Paragraph not found: {exact_text}")


def build_document() -> None:
    if not DIAGRAM.exists():
        raise FileNotFoundError(DIAGRAM)

    document = Document(SOURCE)

    figure1_index, figure1_caption = find_paragraph(document, "Figure 1: System Architecture Diagram")
    replace_picture(document.paragraphs[figure1_index - 1], DIAGRAM)
    replace_text(
        figure1_caption,
        "Figure 1: Updated System Architecture and Output Flow",
        centered=True,
        bold=True,
    )
    replace_text(
        document.paragraphs[figure1_index + 1],
        (
            "Figure 1 presents the implemented request path. A signed-in user submits a dataset or "
            "question through the Next.js interface. The frontend sends the request to FastAPI over "
            "HTTPS with the Supabase JWT. FastAPI validates and routes the request to DatasetAgent and "
            "AnalystAgent. The analytics engine performs the calculations and produces verified insights, "
            "charts, or downloadable reports for the final output box. A separate bidirectional connection "
            "allows FastAPI to read and write account state, metadata, datasets, and reports in Supabase."
        ),
    )

    figure19_index, figure19_caption = find_paragraph(document, "Figure 19: Architecture Diagram")
    replace_picture(document.paragraphs[figure19_index - 1], DIAGRAM)
    replace_text(
        figure19_caption,
        "Figure 19: Updated DataVerse AI Architecture",
        centered=True,
        bold=True,
    )

    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("In the implemented architecture, is layered and modular"):
            replace_text(
                paragraph,
                (
                    "The implementation uses a small two-agent path. After the signed-in user submits a "
                    "file or question, Next.js forwards the request to FastAPI. FastAPI checks the request "
                    "and coordinates DatasetAgent and AnalystAgent. The agents call Pandas, scikit-learn, "
                    "SHAP, and the report builder for deterministic computation. FastAPI uses the separate "
                    "Supabase connection for authentication, metadata, datasets, and saved reports. The "
                    "verified charts and reports are then shown in the user output stage."
                ),
            )
            break

    document.core_properties.title = "DataVerse AI Final Report - Latest Architecture 2026"
    document.core_properties.subject = "Updated one-way architecture and user output flow"
    document.save(OUTPUT)
    print(f"updated_document={OUTPUT}")


if __name__ == "__main__":
    build_document()
