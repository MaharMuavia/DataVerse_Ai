from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "pdf" / "DataVerse_AI_Final_Report_Revised_2026.docx"
OUTPUT = ROOT / "output" / "pdf" / "DataVerse_AI_Final_Report_Simple_Architecture_Formatted_2026.docx"
FONT_NAME = "Times New Roman"
CAPTION_PATTERN = re.compile(r"^(Figure|Table)\s+\d+(?:\.\d+)?:", re.IGNORECASE)


def set_font(target, size: float | None = None) -> None:
    target.font.name = FONT_NAME
    if size is not None:
        target.font.size = Pt(size)
    r_pr = target._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), FONT_NAME)


def format_style(style, size: float, *, line_spacing: float | None = None, justify: bool = False) -> None:
    set_font(style, size)
    if line_spacing is not None:
        style.paragraph_format.line_spacing = line_spacing
    if justify:
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def apply_typography() -> None:
    doc = Document(SOURCE)
    styles = doc.styles

    format_style(styles["Normal"], 12, line_spacing=1.5, justify=True)
    format_style(styles["Heading 1"], 16)
    format_style(styles["Heading 2"], 14)
    format_style(styles["Heading 3"], 13)
    format_style(styles["Title"], 16)
    format_style(styles["List Bullet"], 12, line_spacing=1.5, justify=True)
    format_style(styles["List Number"], 12, line_spacing=1.5, justify=True)
    for name in ("toc 1", "toc 2", "toc 3"):
        if name in styles:
            format_style(styles[name], 12, line_spacing=1.0)

    heading_sizes = {"Heading 1": 16, "Heading 2": 14, "Heading 3": 13}
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name

        if style_name in heading_sizes:
            for run in paragraph.runs:
                set_font(run, heading_sizes[style_name])
                run.bold = True
            continue

        if style_name.lower().startswith("toc "):
            for run in paragraph.runs:
                set_font(run, 12)
            paragraph.paragraph_format.line_spacing = 1.0
            continue

        if index == 1:  # Cover title
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_font(run, 16)
                run.bold = True
            continue

        if 0 <= index <= 11:  # Remaining cover-page text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1.5
            for run in paragraph.runs:
                set_font(run, 12)
            continue

        if CAPTION_PATTERN.match(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1.5
            for run in paragraph.runs:
                set_font(run, 12)
                run.bold = True
            continue

        if style_name in {"Normal", "List Bullet", "List Number"} and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing = 1.5
            for run in paragraph.runs:
                set_font(run, 12)

    # Table text uses the requested typeface while retaining compact sizes so
    # wide academic tables do not overflow the page margins.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        current_size = run.font.size.pt if run.font.size else 10
                        set_font(run, current_size)

    for section in doc.sections:
        for area in (section.header, section.footer):
            for paragraph in area.paragraphs:
                for run in paragraph.runs:
                    set_font(run, run.font.size.pt if run.font.size else 10)

    doc.core_properties.title = "DataVerse AI Final Report - Formatted 2026"
    doc.core_properties.subject = "Times New Roman academic formatting with 16/14/13 point heading hierarchy"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"formatted_report={OUTPUT}")


if __name__ == "__main__":
    apply_typography()
